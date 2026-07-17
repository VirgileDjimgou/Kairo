"""Unit tests for document parsers, chunking, and indexing payloads."""

from io import BytesIO

import pytest
from fakes import FakeEmbeddingProvider, FakeVectorStoreProvider

from app.modules.documents.metadata import (
    classify_archive_access,
    infer_document_language,
    infer_document_language_from_upload,
)
from app.modules.indexing.service import build_chunk_payload
from app.modules.ingestion.chunking import chunk_text, estimate_token_count
from app.providers.parsers import parse_document_bytes
from app.providers.parsers.pdf_docx import parse_docx_bytes, parse_pdf_bytes
from app.providers.parsers.whatsapp import parse_whatsapp_export


def test_parse_txt_bytes() -> None:
    text = parse_document_bytes(b"Hello tenant policy", "notes.txt")
    assert text == "Hello tenant policy"


def test_parse_markdown_bytes() -> None:
    text = parse_document_bytes(b"# Title\n\nBody", "guide.md")
    assert "Title" in text


def test_parse_csv_as_text() -> None:
    text = parse_document_bytes(b"name,role\nAlice,admin", "members.csv")
    assert "Alice" in text


def test_parse_pdf_bytes() -> None:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "tenant policy in pdf")
    pdf_bytes = doc.tobytes()
    doc.close()

    text = parse_pdf_bytes(pdf_bytes)
    assert "tenant policy" in text.lower()


def test_parse_docx_bytes() -> None:
    from docx import Document

    document = Document()
    document.add_paragraph("tenant policy in docx")
    buffer = __import__("io").BytesIO()
    document.save(buffer)

    text = parse_docx_bytes(buffer.getvalue())
    assert "tenant policy" in text.lower()


def test_parse_whatsapp_export() -> None:
    raw = (
        "12/01/2024, 10:30 - Alice: Hello team\n"
        "12/01/2024, 10:31 - Bob: Hi Alice"
    )
    text = parse_whatsapp_export(raw.encode())
    assert "Alice: Hello team" in text
    assert "Bob: Hi Alice" in text


def test_parse_image_bytes_uses_ocr() -> None:
    from PIL import Image, ImageDraw, ImageFont

    image = Image.new("RGB", (900, 240), "white")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("arial.ttf", 64)
    except OSError:
        font = ImageFont.load_default()
    draw.text((40, 70), "HELLO OCR", fill="black", font=font)
    buffer = BytesIO()
    image.save(buffer, format="PNG")

    text = parse_document_bytes(buffer.getvalue(), "scan.png")
    assert "HELLO" in text.upper()


def test_unsupported_extension_raises() -> None:
    with pytest.raises(ValueError, match="not supported yet"):
        parse_document_bytes(b"data", "archive.zip")


def test_chunk_text_splits_long_content() -> None:
    content = "word " * 500
    chunks = chunk_text(content.strip(), chunk_size=100, chunk_overlap=20)
    assert len(chunks) > 1
    assert all(len(chunk) <= 100 for chunk in chunks)


def test_estimate_token_count_is_positive() -> None:
    assert estimate_token_count("one two three") == 3


def test_build_chunk_payload_includes_tenant_and_scope() -> None:
    from uuid import uuid4

    from app.modules.documents.models import Document, DocumentChunk

    tenant_id = uuid4()
    document_id = uuid4()
    version_id = uuid4()
    chunk_id = uuid4()

    document = Document(
        id=document_id,
        tenant_id=tenant_id,
        title="Policy",
        source_type="upload",
        language="fr",
        access_scope="members_only",
        owner_user_id=uuid4(),
        status="uploaded",
    )
    chunk = DocumentChunk(
        id=chunk_id,
        tenant_id=tenant_id,
        document_id=document_id,
        document_version_id=version_id,
        chunk_index=0,
        text="sample",
        language="fr",
    )

    payload = build_chunk_payload(document=document, chunk=chunk)
    assert payload["tenant_id"] == str(tenant_id)
    assert payload["access_scope"] == "members_only"
    assert payload["chunk_id"] == str(chunk_id)
    assert payload["document_version_id"] == str(version_id)


def test_infer_document_language_detects_french_from_text_upload() -> None:
    language = infer_document_language_from_upload(
        file_name="cotisations.txt",
        title="Cotisations 2026",
        file_bytes=b"Bonjour association cotisation annuelle",
    )
    assert language == "fr"


def test_infer_document_language_detects_german_archive_filename() -> None:
    language = infer_document_language(
        file_name="Vorstand_protokoll_de.pdf",
        title="Vorstand Protokoll",
    )
    assert language == "de"


def test_infer_document_language_returns_und_when_ambiguous() -> None:
    language = infer_document_language(
        file_name="notes.txt",
        title="General notes",
        text_sample="No clear linguistic markers here.",
    )
    assert language == "und"


def test_classify_archive_access_is_conservative_for_sensitive_documents() -> None:
    roles = {
        "treasurer": "role-treasurer",
        "auditor": "role-auditor",
        "president": "role-president",
        "vice_president": "role-vice",
        "principal_admin": "role-admin",
        "secretary_general": "role-secretary",
        "censor": "role-censor",
    }

    finance_scope, finance_roles = classify_archive_access("Budget_audit_2026.xlsx", roles)
    governance_scope, governance_roles = classify_archive_access(
        "Proces-verbal_confidential_board_minutes.pdf",
        roles,
    )
    discipline_scope, discipline_roles = classify_archive_access(
        "disciplinary_warning_member.pdf",
        roles,
    )

    assert finance_scope == "role_restricted"
    assert finance_roles == [
        "role-treasurer",
        "role-auditor",
        "role-president",
        "role-vice",
        "role-admin",
    ]
    assert governance_scope == "role_restricted"
    assert governance_roles == [
        "role-secretary",
        "role-president",
        "role-vice",
        "role-admin",
    ]
    assert discipline_scope == "role_restricted"
    assert discipline_roles == [
        "role-censor",
        "role-president",
        "role-vice",
        "role-admin",
    ]


@pytest.mark.asyncio
async def test_fake_indexing_stores_vectors_with_metadata() -> None:
    from uuid import uuid4

    from app.core.config import settings
    from app.modules.documents.models import Document, DocumentChunk
    from app.modules.indexing.service import IndexingService

    previous = settings.indexing_auto_enabled
    settings.indexing_auto_enabled = True
    try:
        tenant_id = uuid4()
        document_id = uuid4()
        version_id = uuid4()
        document = Document(
            id=document_id,
            tenant_id=tenant_id,
            title="Policy",
            source_type="upload",
            language="en",
            access_scope="tenant_public",
            status="uploaded",
        )
        chunk = DocumentChunk(
            id=uuid4(),
            tenant_id=tenant_id,
            document_id=document_id,
            document_version_id=version_id,
            chunk_index=0,
            text="tenant policy chunk",
            language="en",
        )
        vector_store = FakeVectorStoreProvider()
        service = IndexingService(FakeEmbeddingProvider(), vector_store)
        indexed = await service.index_chunks(
            document=document,
            document_version_id=version_id,
            chunks=[chunk],
        )
        assert indexed == 1
        assert len(vector_store.points) == 1
        _, payload = next(iter(vector_store.points.values()))
        assert payload["tenant_id"] == str(tenant_id)
        assert payload["access_scope"] == "tenant_public"
    finally:
        settings.indexing_auto_enabled = previous
